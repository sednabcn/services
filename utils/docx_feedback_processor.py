#!/usr/bin/env python3
"""
DOCX Template Feedback Processor
Automatically injects feedback email into Word document templates
"""
import os
import sys
import argparse
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


class DOCXFeedbackProcessor:
    def __init__(self, feedback_email="feedback@modelphysmat.com"):
        self.feedback_email = feedback_email
        self.processed_files = []
        self.skipped_files = []
        self.failed_files = []
        
    def backup_file(self, file_path):
        """Create backup of original file"""
        backup_path = file_path.with_suffix(f'.backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
        try:
            shutil.copy2(file_path, backup_path)
            return backup_path
        except Exception as e:
            print(f"Warning: Could not create backup for {file_path}: {e}")
            return None
    
    def check_feedback_exists(self, doc):
        """Check if feedback email already exists in document"""
        for paragraph in doc.paragraphs:
            if self.feedback_email.lower() in paragraph.text.lower():
                return True
        return False
    
    def add_feedback_header(self, doc):
        """Add feedback email as header"""
        section = doc.sections[0]
        header = section.header
        
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = header_para.add_run(f"📧 Feedback: {self.feedback_email}")
        run.font.size = 10
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    def add_feedback_footer(self, doc):
        """Add feedback email as footer (default style)"""
        # Add separator line
        separator_para = doc.add_paragraph()
        separator_run = separator_para.add_run("─" * 60)
        separator_run.font.color.rgb = RGBColor(200, 200, 200)
        separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add feedback text
        feedback_para = doc.add_paragraph()
        feedback_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Icon and main text
        icon_run = feedback_para.add_run("📧 ")
        icon_run.font.size = 12
        
        text_run = feedback_para.add_run("We value your feedback! Share your thoughts: ")
        text_run.font.bold = True
        text_run.font.size = 10
        
        email_run = feedback_para.add_run(self.feedback_email)
        email_run.font.bold = True
        email_run.font.color.rgb = RGBColor(0, 100, 200)
        email_run.font.size = 10
        
        # Add improvement text
        improve_para = doc.add_paragraph()
        improve_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        improve_run = improve_para.add_run("🔄 Help us improve - Your input helps us create better communications.")
        improve_run.font.italic = True
        improve_run.font.size = 9
        improve_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def add_feedback_callout(self, doc):
        """Add feedback email as callout box"""
        callout_para = doc.add_paragraph()
        callout_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Create bordered text effect (simulated with characters)
        border_run = callout_para.add_run("┌" + "─" * 50 + "┐\n")
        border_run.font.color.rgb = RGBColor(100, 100, 100)
        
        content_run = callout_para.add_run("│ 📧 FEEDBACK: Send thoughts to ")
        content_run.font.size = 10
        
        email_run = callout_para.add_run(self.feedback_email)
        email_run.font.bold = True
        email_run.font.color.rgb = RGBColor(0, 100, 200)
        email_run.font.size = 10
        
        space_run = callout_para.add_run(" │\n")
        space_run.font.color.rgb = RGBColor(100, 100, 100)
        
        bottom_run = callout_para.add_run("└" + "─" * 50 + "┘")
        bottom_run.font.color.rgb = RGBColor(100, 100, 100)
    
    def process_template(self, template_path, style="footer", create_backup=True):
        """Process a single DOCX template"""
        try:
            # Create backup if requested
            if create_backup:
                backup_path = self.backup_file(template_path)
                if backup_path:
                    print(f"  📄 Backup created: {backup_path.name}")
            
            # Open document
            doc = Document(template_path)
            
            # Check if feedback already exists
            if self.check_feedback_exists(doc):
                print(f"  ⚠️  Feedback email already exists, skipping...")
                self.skipped_files.append({
                    'file': str(template_path),
                    'reason': 'Feedback already exists'
                })
                return False
            
            # Add feedback based on style
            if style == "header":
                self.add_feedback_header(doc)
            elif style == "footer":
                self.add_feedback_footer(doc)
            elif style == "callout":
                self.add_feedback_callout(doc)
            else:
                raise ValueError(f"Unknown style: {style}")
            
            # Save document
            doc.save(template_path)
            
            self.processed_files.append({
                'file': str(template_path),
                'style': style,
                'feedback_email': self.feedback_email
            })
            
            print(f"  ✅ Processed with {style} style")
            return True
            
        except Exception as e:
            print(f"  ❌ Error processing: {e}")
            self.failed_files.append({
                'file': str(template_path),
                'error': str(e)
            })
            return False
    
    def process_domain_templates(self, templates_dir, style="footer", create_backup=True):
        """Process all templates in domain-based structure"""
        templates_path = Path(templates_dir)
        
        if not templates_path.exists():
            print(f"Templates directory not found: {templates_dir}")
            return
        
        print(f"Processing templates in: {templates_path}")
        
        # Process each domain directory
        for domain_dir in templates_path.iterdir():
            if domain_dir.is_dir():
                print(f"\n📁 Processing domain: {domain_dir.name}")
                
                # Find all DOCX files in domain
                docx_files = list(domain_dir.glob("*.docx"))
                
                if not docx_files:
                    print(f"  No DOCX files found in {domain_dir.name}")
                    continue
                
                for docx_file in docx_files:
                    print(f"  📄 Processing: {docx_file.name}")
                    self.process_template(docx_file, style, create_backup)
    
    def generate_processing_report(self, output_path="docx_feedback_processing_report.md"):
        """Generate processing report"""
        report_content = f"""# DOCX Feedback Processing Report

**Processing Time:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**Feedback Email:** `{self.feedback_email}`

## Summary

Skipped:** {len(self.skipped_files)} files
- **Failed:** {len(self.failed_files)} files

## Processed Files

"""
        
        for file_info in self.processed_files:
            report_content += f"### ✅ {Path(file_info['file']).name}\n"
            report_content += f"- **Path:** `{file_info['file']}`\n"
            report_content += f"- **Style:** {file_info['style']}\n"
            report_content += f"- **Feedback Email:** `{file_info['feedback_email']}`\n\n"
        
        if self.skipped_files:
            report_content += "## Skipped Files\n\n"
            for file_info in self.skipped_files:
                report_content += f"### ⚠️ {Path(file_info['file']).name}\n"
                report_content += f"- **Reason:** {file_info['reason']}\n\n"
        
        if self.failed_files:
            report_content += "## Failed Files\n\n"
            for file_info in self.failed_files:
                report_content += f"### ❌ {Path(file_info['file']).name}\n"
                report_content += f"- **Error:** {file_info['error']}\n\n"
        
        report_content += f"""
## Feedback Email Integration

The feedback email `{self.feedback_email}` has been injected into all successfully processed templates. Recipients can use this email to:

- Provide feedback on the communication
- Share improvement suggestions
- Report any concerns
- Express interest in partnerships

All feedback will be automatically tracked and forwarded to the appropriate team members.
"""
        
        # Write report to file
        with open(output_path, 'w') as f:
            f.write(report_content)
        
        print(f"📋 Processing report generated: {output_path}")
        return output_path


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Process DOCX templates with feedback email")
    parser.add_argument("--templates-dir", default="campaign-templates", 
                       help="Templates directory")
    parser.add_argument("--feedback-email", default="feedback@modelphysmat.com",
                       help="Feedback email address")
    parser.add_argument("--style", choices=["header", "footer", "callout"],
                       default="footer", help="Feedback injection style")
    parser.add_argument("--create-backups", action="store_true",
                       help="Create backups of original files")
    parser.add_argument("--create-report", action="store_true",
                       help="Generate processing report")
    parser.add_argument("--domain", help="Process specific domain only")
    
    args = parser.parse_args()
    
    processor = DOCXFeedbackProcessor(args.feedback_email)
    
    print(f"📧 DOCX Feedback Processor")
    print(f"📁 Templates Dir: {args.templates_dir}")
    print(f"📧 Feedback Email: {args.feedback_email}")
    print(f"🎨 Style: {args.style}")
    print(f"💾 Create Backups: {args.create_backups}")
    print("=" * 60)
    
    if args.domain:
        # Process specific domain
        domain_path = Path(args.templates_dir) / args.domain
        if domain_path.exists():
            print(f"Processing specific domain: {args.domain}")
            for docx_file in domain_path.glob("*.docx"):
                print(f"📄 Processing: {docx_file.name}")
                processor.process_template(docx_file, args.style, args.create_backups)
        else:
            print(f"Domain directory not found: {domain_path}")
    else:
        # Process all domains
        processor.process_domain_templates(args.templates_dir, args.style, args.create_backups)
    
    # Generate report if requested
    if args.create_report:
        processor.generate_processing_report()
    
    print(f"\n📊 PROCESSING SUMMARY:")
    print(f"✅ Successfully processed: {len(processor.processed_files)} files")
    print(f"⚠️ Skipped: {len(processor.skipped_files)} files")
    print(f"❌ Failed: {len(processor.failed_files)} files")
    
    if processor.processed_files:
        print(f"\n📧 Feedback email '{args.feedback_email}' has been added to all processed templates")


if __name__ == "__main__":
    main()Processed Successfully:** {len(self.processed_files)} files
- **
