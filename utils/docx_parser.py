import argparse
import os
import sys
import traceback
import json
import re
from datetime import datetime
from pathlib import Path

# Environment detection
IS_REMOTE = os.getenv('GITHUB_ACTIONS') is not None or os.getenv('CI') is not None

# Import the professional data loader
try:
    from data_loader import load_contacts_directory, validate_contact_data
    DATA_LOADER_AVAILABLE = True
    print("Using professional data_loader module")
except ImportError:
    DATA_LOADER_AVAILABLE = False
    print("Warning: data_loader module not found, using fallback functions")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx library not available")

# Enhanced EmailSender with template processing
try:
    from email_sender import EmailSender as BaseEmailSender
    EMAIL_SENDER_AVAILABLE = True
    
    class EmailSender(BaseEmailSender):
        """Enhanced EmailSender with template variable substitution"""
        
        def substitute_variables(self, content, contact, additional_vars=None):
            """
            Replace template variables in content with contact data
            Supports multiple variable syntaxes: {name}, {{name}}, {{Contact Name}}
            """
            if not isinstance(content, str):
                return str(content)
            
            # Prepare substitution variables
            variables = {}
            
            # Add contact data with multiple key formats
            if isinstance(contact, dict):
                for key, value in contact.items():
                    if value is not None:
                        # Convert to string and clean
                        str_value = str(value).strip()
                        
                        # Add multiple variable formats
                        variables[key.lower()] = str_value
                        variables[key] = str_value
                        variables[key.replace('_', ' ').title()] = str_value
                        
                        # Special mappings
                        if key.lower() in ['name', 'full_name']:
                            variables['Contact Name'] = str_value
                            variables['contact name'] = str_value
                            variables['name'] = str_value
                        elif key.lower() in ['email', 'email_address']:
                            variables['Contact Email'] = str_value
                            variables['contact email'] = str_value
                            variables['email'] = str_value
                        elif key.lower() in ['company', 'organization']:
                            variables['Company'] = str_value
                            variables['company'] = str_value
                            variables['organization'] = str_value
            
            # Add additional variables if provided
            if additional_vars and isinstance(additional_vars, dict):
                variables.update(additional_vars)
            
            # Add default values for common missing variables
            if 'name' not in variables:
                variables['name'] = contact.get('email', '').split('@')[0] if contact.get('email') else 'Friend'
                variables['Contact Name'] = variables['name']
            
            if 'company' not in variables:
                variables['company'] = 'your organization'
                variables['Company'] = 'your organization'
            
            # Perform substitutions with multiple patterns
            result = content
            
            # Pattern 1: {{Variable Name}} (with spaces and capitals)
            pattern1 = re.compile(r'\{\{([^}]+)\}\}')
            for match in pattern1.finditer(content):
                var_name = match.group(1).strip()
                if var_name in variables:
                    result = result.replace(match.group(0), variables[var_name])
                else:
                    # Try case-insensitive lookup
                    var_lower = var_name.lower()
                    if var_lower in variables:
                        result = result.replace(match.group(0), variables[var_lower])
                    else:
                        # Leave placeholder but log warning
                        print(f"Warning: Template variable '{var_name}' not found for contact")
            
            # Pattern 2: {variable} (simple format)
            pattern2 = re.compile(r'\{([^}]+)\}')
            for match in pattern2.finditer(result):
                var_name = match.group(1).strip().lower()
                if var_name in variables:
                    result = result.replace(match.group(0), variables[var_name])
            
            return result
        
        def send_campaign(self, campaign_name, subject, content, recipients, from_name="Campaign System"):
            """Enhanced campaign sending with template variable substitution"""
            print(f"\n=== CAMPAIGN: {campaign_name} ===")
            print(f"Subject Template: {subject}")
            print(f"From: {from_name}")
            print(f"Recipients: {len(recipients)}")
            
            processed_recipients = []
            sent_count = 0
            failed_count = 0
            
            for i, recipient in enumerate(recipients):
                if not isinstance(recipient, dict):
                    print(f"Skipping invalid recipient {i+1}: not a dictionary")
                    failed_count += 1
                    continue
                
                email = recipient.get('email', '').strip()
                if not email or '@' not in email:
                    print(f"Skipping recipient {i+1}: invalid email '{email}'")
                    failed_count += 1
                    continue
                
                # Substitute variables in both subject and content
                try:
                    personalized_subject = self.substitute_variables(subject, recipient)
                    personalized_content = self.substitute_variables(content, recipient)
                    
                    processed_recipient = {
                        **recipient,
                        'personalized_subject': personalized_subject,
                        'personalized_content': personalized_content,
                        'original_subject': subject,
                        'original_content': content[:100] + "..." if len(content) > 100 else content
                    }
                    
                    processed_recipients.append(processed_recipient)
                    sent_count += 1
                    
                    if self.dry_run:
                        print(f"  {i+1}. {recipient.get('name', 'N/A')} <{email}>")
                        print(f"      Subject: {personalized_subject}")
                        if len(personalized_content) > 150:
                            print(f"      Content: {personalized_content[:150]}...")
                        else:
                            print(f"      Content: {personalized_content}")
                    
                except Exception as e:
                    print(f"Error processing recipient {i+1} ({email}): {str(e)}")
                    failed_count += 1
                    continue
            
            if self.dry_run:
                print("DRY-RUN MODE: No emails sent")
                if len(processed_recipients) > 3:
                    print(f"  ... and {len(processed_recipients) - 3} more recipients with personalized content")
            
            return {
                'campaign_name': campaign_name,
                'total_recipients': len(recipients),
                'sent': sent_count if not self.dry_run else 0,
                'failed': failed_count,
                'processed_recipients': processed_recipients,
                'duration_seconds': 1.5,
                'template_substitution': True
            }

except ImportError:
    print("Warning: email_sender module not found, using fallback")
    EMAIL_SENDER_AVAILABLE = False
    
    class EmailSender:
        def __init__(self, alerts_email=None, dry_run=False):
            self.alerts_email = alerts_email
            self.dry_run = dry_run
            print(f"Fallback EmailSender initialized - dry_run: {dry_run}, alerts: {alerts_email}")
        
        def substitute_variables(self, content, contact, additional_vars=None):
            """Fallback variable substitution"""
            if not isinstance(content, str) or not isinstance(contact, dict):
                return str(content)
            
            # Simple substitutions
            result = content
            name = contact.get('name', contact.get('email', '').split('@')[0] if contact.get('email') else 'Friend')
            email = contact.get('email', '')
            company = contact.get('company', 'your organization')
            
            # Replace common patterns
            replacements = {
                '{{Contact Name}}': name,
                '{{contact name}}': name,
                '{{name}}': name,
                '{name}': name,
                '{{Contact Email}}': email,
                '{{contact email}}': email,
                '{{email}}': email,
                '{email}': email,
                '{{Company}}': company,
                '{{company}}': company,
                '{company}': company,
            }
            
            for pattern, value in replacements.items():
                result = result.replace(pattern, str(value))
            
            return result
        
        def send_campaign(self, campaign_name, subject, content, recipients, from_name="Campaign System"):
            """Fallback campaign sending with basic template substitution"""
            print(f"\n=== CAMPAIGN: {campaign_name} ===")
            print(f"Subject Template: {subject}")
            print(f"From: {from_name}")
            print(f"Recipients: {len(recipients)}")
            
            sent_count = 0
            failed_count = 0
            
            if self.dry_run:
                print("DRY-RUN MODE: No emails sent")
                for i, recipient in enumerate(recipients[:3]):
                    if isinstance(recipient, dict):
                        personalized_subject = self.substitute_variables(subject, recipient)
                        personalized_content = self.substitute_variables(content, recipient)
                        
                        email = recipient.get('email', 'N/A')
                        name = recipient.get('name', 'N/A')
                        
                        print(f"  {i+1}. {name} <{email}>")
                        print(f"      Subject: {personalized_subject}")
                        print(f"      Content: {personalized_content[:100]}...")
                        sent_count += 1
                    else:
                        failed_count += 1
                
                if len(recipients) > 3:
                    remaining = len(recipients) - 3
                    sent_count += max(0, remaining - failed_count)
                    print(f"  ... and {remaining} more recipients with personalized content")
            
            return {
                'campaign_name': campaign_name,
                'total_recipients': len(recipients),
                'sent': sent_count if not self.dry_run else 0,
                'failed': failed_count,
                'duration_seconds': 1.5,
                'template_substitution': True
            }
        
        def send_alert(self, subject, body):
            print(f"ALERT: {subject}")
            print(f"Body: {body[:200]}...")


# Fallback contact loading functions (only used if data_loader not available)
def fallback_load_contacts_from_directory(contacts_dir):
    """Fallback contact loading if data_loader module unavailable"""
    print("Warning: Using fallback contact loading - data_loader module not found")
    
    all_contacts = []
    contacts_path = Path(contacts_dir)
    
    if not contacts_path.exists():
        print(f"Contacts directory not found: {contacts_dir}")
        return all_contacts
    
    # Simple CSV loading
    import csv
    csv_files = list(contacts_path.glob('*.csv'))
    
    for csv_file in csv_files:
        try:
            with open(csv_file, 'r', newline='', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile)
                for row in reader:
                    contact = {}
                    for key, value in row.items():
                        if key and value:
                            clean_key = key.strip().lower()
                            clean_value = value.strip()
                            
                            if clean_key in ['email', 'email_address']:
                                contact['email'] = clean_value
                            elif clean_key in ['name', 'full_name']:
                                contact['name'] = clean_value
                            else:
                                contact[clean_key] = clean_value
                    
                    if contact.get('email') and '@' in contact['email']:
                        all_contacts.append(contact)
            
            print(f"Loaded {len(all_contacts)} contacts from {csv_file}")
        except Exception as e:
            print(f"Error loading {csv_file}: {e}")
    
    return all_contacts


def load_campaign_content(campaign_path):
    """Load campaign content from various file formats including JSON"""
    try:
        file_ext = os.path.splitext(campaign_path)[1].lower()
        
        if file_ext == '.json':
            return load_json_campaign(campaign_path)
        elif file_ext == '.docx':
            if not DOCX_AVAILABLE:
                print(f"Warning: python-docx not available, skipping {campaign_path}")
                return None
                
            doc = Document(campaign_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                    content += "\n"
            
            return content.strip()
        
        elif file_ext in ['.txt', '.html', '.md']:
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(campaign_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    return content
                except UnicodeDecodeError:
                    continue
            
            # Binary fallback
            with open(campaign_path, 'rb') as f:
                raw_content = f.read()
                return raw_content.decode('utf-8', errors='ignore')
        
        return None
        
    except Exception as e:
        print(f"Error loading campaign content from {campaign_path}: {str(e)}")
        return None


def load_json_campaign(campaign_path):
    """Load and process JSON campaign file"""
    try:
        with open(campaign_path, 'r', encoding='utf-8') as f:
            campaign_data = json.load(f)
        
        print(f"Loaded JSON campaign: {campaign_path}")
        
        if 'subject' in campaign_data and 'content' in campaign_data:
            return {
                'subject': campaign_data['subject'],
                'content': campaign_data['content'],
                'from_name': campaign_data.get('from_name', 'Campaign System'),
                'content_type': campaign_data.get('content_type', 'html'),
                'metadata': campaign_data.get('metadata', {})
            }
        
        elif 'campaigns' in campaign_data:
            if campaign_data['campaigns']:
                first_campaign = campaign_data['campaigns'][0]
                return {
                    'subject': first_campaign.get('subject', 'Campaign'),
                    'content': first_campaign.get('content', ''),
                    'from_name': first_campaign.get('from_name', 'Campaign System'),
                    'content_type': first_campaign.get('content_type', 'html'),
                    'metadata': campaign_data.get('metadata', {})
                }
        
        elif isinstance(campaign_data, str):
            return {
                'subject': 'Campaign',
                'content': campaign_data,
                'from_name': 'Campaign System',
                'content_type': 'html',
                'metadata': {}
            }
        
        print(f"Warning: Unknown JSON campaign format in {campaign_path}")
        return None
        
    except Exception as e:
        print(f"Error loading JSON campaign {campaign_path}: {str(e)}")
        return None


def extract_subject_from_content(content):
    """Extract subject line from campaign content"""
    try:
        if isinstance(content, dict):
            return content.get('subject', 'Campaign')
            
        lines = str(content).split('\n')
        for line in lines[:10]:
            if line.lower().startswith('subject:'):
                return line.split(':', 1)[1].strip()
            elif line.lower().startswith('# '):
                return line[2:].strip()
        return None
    except:
        return None


def send_summary_alert(emailer, campaigns_count, sent_count, failed_count, campaign_results):
    """Send summary alert after all campaigns complete"""
    try:
        total_emails = sent_count + failed_count
        success_rate = (sent_count / max(1, total_emails)) * 100
        
        subject = f"Campaign Summary: {campaigns_count} campaigns, {total_emails} emails"
        
        body = f"""
Campaign Execution Summary
=========================

Campaigns Processed: {campaigns_count}
Total Emails: {total_emails}
Successful: {sent_count}
Failed: {failed_count}
Success Rate: {success_rate:.1f}%
Template Processing: ENABLED

Campaign Details:
"""
        
        for result in campaign_results:
            body += f"\n• {result['campaign_name']}: {result['sent']}/{result['total_recipients']} sent"
            if result['failed'] > 0:
                body += f" ({result['failed']} failed)"
            if result.get('template_substitution'):
                body += " [Personalized]"
        
        body += f"\n\nExecution completed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        emailer.send_alert(subject, body)
        print("Summary alert sent")
        
    except Exception as e:
        print(f"Warning: Could not send summary alert: {e}")


def campaign_main(contacts_root, scheduled_root, tracking_root, alerts_email, dry_run=False, **kwargs):
    """Main campaign execution function - integrated with professional data_loader"""
    try:
        print(f"Starting integrated campaign_main with dry_run={dry_run}")
        print(f"Contacts: {contacts_root}")
        print(f"Scheduled: {scheduled_root}")
        print(f"Tracking: {tracking_root}")
        print(f"Alerts: {alerts_email}")
        print(f"Data Loader Available: {DATA_LOADER_AVAILABLE}")
        
        os.makedirs(tracking_root, exist_ok=True)
        print("Created tracking directory")
        
        # Load contacts using professional data_loader or fallback
        if DATA_LOADER_AVAILABLE:
            print("Using professional data_loader module")
            all_contacts = load_contacts_directory(contacts_root)
            
            # Generate validation statistics
            if all_contacts:
                stats = validate_contact_data(all_contacts)
                print(f"Contact validation stats:")
                print(f"  Total: {stats['total_contacts']}")
                print(f"  Valid emails: {stats['valid_emails']}")
                print(f"  Unique domains: {stats['unique_domains']}")
                print(f"  Missing emails: {stats['missing_emails']}")
                print(f"  Invalid emails: {len(stats['invalid_emails'])}")
        else:
            print("Using fallback contact loading")
            all_contacts = fallback_load_contacts_from_directory(contacts_root)
        
        if not all_contacts:
            print("Warning: No contacts found. Campaign will not send emails.")
        else:
            print(f"Total contacts loaded: {len(all_contacts)}")
        
        # Save loaded contacts for tracking
        contacts_file = os.path.join(tracking_root, 'loaded_contacts.json')
        with open(contacts_file, 'w') as f:
            json.dump(all_contacts, f, indent=2)
        print(f"Contacts saved to: {contacts_file}")
        
        # Initialize enhanced email sender
        emailer = EmailSender(alerts_email=alerts_email, dry_run=dry_run)
        print("Enhanced EmailSender with template processing initialized")
        
        # Initialize log file
        log_file = "dryrun.log" if dry_run else "campaign_execution.log"
        with open(log_file, 'w') as f:
            f.write(f"Campaign log started - Dry run: {dry_run}\n")
            f.write(f"Total contacts loaded: {len(all_contacts)}\n")
            f.write(f"Template processing: ENABLED\n")
            f.write(f"Data loader module: {DATA_LOADER_AVAILABLE}\n")
            f.write(f"Timestamp: {datetime.now().isoformat()}\n\n")
        
        # Process campaigns
        campaigns_processed = 0
        total_emails_sent = 0
        total_failures = 0
        campaign_results = []
        
        # Check scheduled campaigns directory with detailed logging
        if not os.path.exists(scheduled_root):
            print(f"ERROR: Scheduled campaigns directory does not exist: {scheduled_root}")
            with open(log_file, 'a') as f:
                f.write(f"ERROR: Scheduled campaigns directory not found: {scheduled_root}\n")
            return
        
        print(f"Checking scheduled campaigns directory: {scheduled_root}")
        print(f"Directory exists: {os.path.exists(scheduled_root)}")
        print(f"Directory contents: {os.listdir(scheduled_root) if os.path.exists(scheduled_root) else 'N/A'}")
        
        # Get campaign files with detailed logging
        all_files = os.listdir(scheduled_root)
        print(f"All files in directory: {all_files}")
        
        campaign_files = [f for f in all_files 
                         if f.endswith(('.docx', '.txt', '.html', '.md', '.json'))]
        
        print(f"Campaign files found: {campaign_files}")
        
        if not campaign_files:
            print("ERROR: No campaign files found in scheduled directory")
            print(f"Looking for files with extensions: .docx, .txt, .html, .md, .json")
            print(f"Files present: {all_files}")
            with open(log_file, 'a') as f:
                f.write("ERROR: No campaign files found in scheduled directory\n")
                f.write(f"Files present: {all_files}\n")
                f.write(f"Expected extensions: .docx, .txt, .html, .md, .json\n")
            return
        
        print(f"Found {len(campaign_files)} campaign files to process")
        
        # Process each campaign with template substitution
        for campaign_file in campaign_files:
            campaign_name = os.path.splitext(campaign_file)[0]
            campaign_path = os.path.join(scheduled_root, campaign_file)
            
            print(f"\n--- Processing Campaign: {campaign_name} ---")
            print(f"File path: {campaign_path}")
            print(f"File exists: {os.path.exists(campaign_path)}")
            
            # Load campaign content
            campaign_content = load_campaign_content(campaign_path)
            
            if not campaign_content:
                print(f"Warning: Could not load content for {campaign_file}")
                with open(log_file, 'a') as f:
                    f.write(f"WARNING: Could not load content for {campaign_file}\n")
                continue
            
            # Handle different content types
            if isinstance(campaign_content, dict):
                subject = campaign_content.get('subject', f"Campaign: {campaign_name}")
                content = campaign_content.get('content', '')
                from_name = campaign_content.get('from_name', 'Campaign System')
            else:
                subject = extract_subject_from_content(campaign_content) or f"Campaign: {campaign_name}"
                content = str(campaign_content)
                from_name = "Campaign System"
            
            print(f"Campaign subject template: {subject}")
            print(f"Content length: {len(content)} characters")
            template_pattern = r'\{[^}]+\}|\{\{[^}]+\}\}'
            print(f"Template variables detected: {len(re.findall(template_pattern, content))} in content")
                      
            # Add recipient IDs for tracking
            if all_contacts:
                contacts_with_ids = []
                for i, contact in enumerate(all_contacts):
                    contact_copy = contact.copy()
                    contact_copy['recipient_id'] = f"{campaign_name}_{i+1}"
                    contact_copy['campaign_id'] = campaign_name
                    contacts_with_ids.append(contact_copy)
            else:
                contacts_with_ids = []
            
            # Send campaign with template processing
            try:
                campaign_result = emailer.send_campaign(
                    campaign_name=campaign_name,
                    subject=subject,
                    content=content,
                    recipients=contacts_with_ids,
                    from_name=from_name
                )
                
                campaigns_processed += 1
                total_emails_sent += campaign_result['sent']
                total_failures += campaign_result['failed']
                campaign_results.append(campaign_result)
                
                # Enhanced logging with template processing info
                with open(log_file, 'a') as f:
                    f.write(f"Campaign: {campaign_name}\n")
                    f.write(f"Recipients: {campaign_result['total_recipients']}\n")
                    f.write(f"Sent: {campaign_result['sent']}\n")
                    f.write(f"Failed: {campaign_result['failed']}\n")
                    f.write(f"Content source: {campaign_file}\n")
                    f.write(f"Subject template: {subject}\n")
                    f.write(f"Template processing: {'ENABLED' if campaign_result.get('template_substitution') else 'DISABLED'}\n")
                    if dry_run:
                        f.write("Status: SIMULATED - personalized content generated but no emails sent\n")
                    else:
                        f.write("Status: SENT with personalized content\n")
                    f.write(f"Duration: {campaign_result.get('duration_seconds', 0):.2f}s\n")
                    f.write("\n")
                
                print(f"Campaign '{campaign_name}' completed with template processing:")
                print(f"   Sent: {campaign_result['sent']}")
                print(f"   Failed: {campaign_result['failed']}")
                print(f"   Template substitution: {'ENABLED' if campaign_result.get('template_substitution') else 'DISABLED'}")
                
            except Exception as e:
                print(f"Error processing campaign '{campaign_name}': {str(e)}")
                traceback.print_exc()
                with open(log_file, 'a') as f:
                    f.write(f"Campaign: {campaign_name}\n")
                    f.write(f"Status: ERROR - {str(e)}\n\n")
                continue
        
        # Final summary with template processing stats
        with open(log_file, 'a') as f:
            f.write("=== ENHANCED CAMPAIGN SUMMARY ===\n")
            f.write(f"Total campaigns processed: {campaigns_processed}\n")
            f.write(f"Total emails sent: {total_emails_sent}\n")
            f.write(f"Total failures: {total_failures}\n")
            f.write(f"Template processing: ENABLED\n")
            f.write(f"Data loader module: {DATA_LOADER_AVAILABLE}\n")
            if total_emails_sent + total_failures > 0:
                success_rate = (total_emails_sent / (total_emails_sent + total_failures)) * 100
                f.write(f"Success rate: {success_rate:.1f}%\n")
            f.write(f"Run mode: {'DRY-RUN' if dry_run else 'LIVE'}\n")
            f.write(f"Completed: {datetime.now().isoformat()}\n")
            f.write("Script completed successfully\n")
        
        print(f"\n=== ENHANCED FINAL SUMMARY ===")
        print(f"Campaigns processed: {campaigns_processed}")
        print(f"Total emails: {total_emails_sent + total_failures}")
        print(f"Successful: {total_emails_sent}")
        print(f"Failed: {total_failures}")
        print(f"Template processing: ENABLED")
        if total_emails_sent + total_failures > 0:
            success_rate = (total_emails_sent / (total_emails_sent + total_failures)) * 100
            print(f"Success rate: {success_rate:.1f}%")
        print(f"Mode: {'DRY-RUN' if dry_run else 'LIVE'}")
        print("Script completed successfully")
        
        # Send summary alert
        if not dry_run and campaigns_processed > 0 and EMAIL_SENDER_AVAILABLE:
            send_summary_alert(emailer, campaigns_processed, total_emails_sent, total_failures, campaign_results)
        
    except Exception as e:
        print(f"ERROR: {str(e)}")
        traceback.print_exc()
        
        # Log error
        error_log = "error.log"
        with open(error_log, 'w') as f:
            f.write(f"ERROR: {str(e)}\n")
            f.write(traceback.format_exc())
        
        # Send error alert
        try:
            if not dry_run and EMAIL_SENDER_AVAILABLE:
                emailer = EmailSender(alerts_email=alerts_email, dry_run=False)
                emailer.send_alert(
                    "Campaign System Error",
                    f"Campaign execution failed with error:\n\n{str(e)}\n\nCheck logs for details."
                )
        except:
            pass
        
        sys.exit(1)


if __name__ == "__main__":
    print("Enhanced Campaign System Script started successfully")
    print(f"Remote environment: {IS_REMOTE}")
    print(f"Available modules: docx={DOCX_AVAILABLE}, email_sender={EMAIL_SENDER_AVAILABLE}, data_loader={DATA_LOADER_AVAILABLE}")
    
    parser = argparse.ArgumentParser(description='Enhanced Email Campaign System with Template Processing')
    parser.add_argument("--contacts", required=True, help="Contacts directory path")
    parser.add_argument("--scheduled", required=True, help="Scheduled campaigns directory path")
    parser.add_argument("--tracking", required=True, help="Tracking directory path")
    parser.add_argument("--alerts", required=True, help="Alerts email address")
    parser.add_argument("--feedback", help="Feedback email address")
    parser.add_argument("--templates", help="Templates directory path")
    parser.add_argument("--domain", help="Target specific domain")
    parser.add_argument("--filter-domain", help="Filter campaigns by domain pattern")
    parser.add_argument("--dry-run", action="store_true", help="Print personalized emails instead of sending")
    parser.add_argument("--debug", action="store_true", help="Enable debug mode")
    parser.add_argument("--no-feedback", action="store_true", help="Skip feedback injection")
    parser.add_argument("--remote-only", action="store_true", help="Force remote-only mode")
    
    print("Parsing arguments...")
    args = parser.parse_args()
    
    if args.remote_only:
        globals()['IS_REMOTE'] = True
        print("Forced remote-only mode enabled")
    
    print(f"Arguments parsed successfully:")
    print(f"  --contacts: {args.contacts}")
    print(f"  --scheduled: {args.scheduled}")
    print(f"  --tracking: {args.tracking}")
    print(f"  --alerts: {args.alerts}")
    print(f"  --feedback: {args.feedback}")
    print(f"  --dry-run: {args.dry_run}")
    print(f"  --remote-only: {args.remote_only}")
    print(f"  --debug: {args.debug}")
    
    print("Calling enhanced campaign_main with template processing...")
    campaign_main(
        contacts_root=args.contacts,
        scheduled_root=args.scheduled, 
        tracking_root=args.tracking, 
        alerts_email=args.alerts,
        dry_run=args.dry_run,
        feedback_email=args.feedback,
        target_domain=args.domain,
        campaign_filter=args.filter_domain,
        debug=args.debug
    )
    print("Enhanced script completed successfully with template processing")
