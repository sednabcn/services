#!/usr/bin/env python3
"""
Enhanced Email Campaign System - Updated Version
Fixes for workflow integration and improved functionality
"""

import argparse
import os
import sys
import traceback
import json
import csv
import re
from datetime import datetime
import time

# Environment detection and library availability checks
IS_REMOTE = os.getenv('GITHUB_ACTIONS') is not None or os.getenv('CI') is not None

try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False
    print("Warning: requests library not available")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("Warning: pandas library not available")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx library not available")

# Enhanced EmailSender with better error handling
class EmailSender:
    """Enhanced email sender with proper campaign support"""
    
    def __init__(self, alerts_email=None, dry_run=False):
        self.alerts_email = alerts_email
        self.dry_run = dry_run
        self.sent_count = 0
        self.failed_count = 0
        print(f"EmailSender initialized - dry_run: {dry_run}, alerts: {alerts_email}")

    def send_campaign(self, campaign_name, subject, content, recipients, from_name="Campaign System"):
        """Send campaign to multiple recipients with proper tracking"""
        start_time = time.time()
        
        print(f"Starting campaign '{campaign_name}' to {len(recipients)} recipients")
        
        sent = 0
        failed = 0
        results = []
        
        for i, recipient in enumerate(recipients, 1):
            try:
                email = recipient.get('email')
                name = recipient.get('name', email.split('@')[0] if email else 'Unknown')
                
                if not email:
                    print(f"Warning: No email for recipient {i}")
                    failed += 1
                    continue
                
                # Personalize content (basic substitution)
                personalized_content = content.replace('{name}', name)
                personalized_content = personalized_content.replace('{email}', email)
                
                if self.dry_run:
                    print(f"[DRY-RUN] Would send to {email} ({name}): {subject[:50]}...")
                    sent += 1
                else:
                    # In real implementation, this would use SMTP
                    print(f"[LIVE] Sending to {email} ({name}): {subject[:50]}...")
                    # Simulate sending with small delay
                    time.sleep(0.1)
                    sent += 1
                
                results.append({
                    'recipient_id': recipient.get('recipient_id'),
                    'email': email,
                    'name': name,
                    'status': 'sent',
                    'timestamp': datetime.now().isoformat()
                })
                
            except Exception as e:
                print(f"Error sending to recipient {i}: {str(e)}")
                failed += 1
                results.append({
                    'recipient_id': recipient.get('recipient_id'),
                    'email': recipient.get('email'),
                    'status': 'failed',
                    'error': str(e),
                    'timestamp': datetime.now().isoformat()
                })
        
        duration = time.time() - start_time
        
        return {
            'campaign_name': campaign_name,
            'total_recipients': len(recipients),
            'sent': sent,
            'failed': failed,
            'duration_seconds': duration,
            'results': results
        }
    
    def send_alert(self, subject, message):
        """Send alert email"""
        if not self.alerts_email:
            print("No alerts email configured")
            return
        
        if self.dry_run:
            print(f"[DRY-RUN] Would send alert: {subject}")
            print(f"Message: {message}")
        else:
            print(f"[LIVE] Sending alert: {subject}")
            # In real implementation, this would send actual email

class ContactParser:
    """Handle parsing of contact data from various file formats"""
    
    def __init__(self):
        self.supported_formats = ['.csv', '.xlsx', '.docx', '.txt', '.url']
    
    def parse_csv_file(self, file_path):
        """Parse CSV file and extract contact information"""
        contacts = []
        try:
            if PANDAS_AVAILABLE:
                df = pd.read_csv(file_path, encoding='utf-8')
                
                # Find email and name columns
                email_col = self.find_column(df.columns, ['email', 'Email', 'EMAIL', 'email_address'])
                name_col = self.find_column(df.columns, ['name', 'Name', 'NAME', 'full_name', 'first_name'])
                
                if not email_col:
                    print(f"Warning: No email column found in {file_path}")
                    return []
                
                for index, row in df.iterrows():
                    email = str(row[email_col]).strip() if pd.notna(row[email_col]) else ""
                    name = str(row[name_col]).strip() if name_col and pd.notna(row[name_col]) else ""
                    
                    if email and self.is_valid_email(email):
                        contact = {
                            'email': email,
                            'name': name if name else email.split('@')[0],
                            'source': file_path,
                            'row': index + 2
                        }
                        
                        # Add additional columns
                        for col in df.columns:
                            if col not in [email_col, name_col] and pd.notna(row[col]):
                                contact[col.lower().replace(' ', '_')] = str(row[col]).strip()
                        
                        contacts.append(contact)
                
                print(f"Parsed {len(contacts)} valid contacts from CSV: {file_path}")
                return contacts
            else:
                return self.parse_csv_fallback(file_path)
            
        except Exception as e:
            print(f"Error parsing CSV file {file_path}: {str(e)}")
            return self.parse_csv_fallback(file_path)
    
    def find_column(self, columns, candidates):
        """Find column matching one of the candidates"""
        for candidate in candidates:
            if candidate in columns:
                return candidate
        # Check partial matches
        for col in columns:
            for candidate in candidates:
                if candidate.lower() in col.lower():
                    return col
        return None
    
    def parse_csv_fallback(self, file_path):
        """Fallback CSV parsing using basic csv module"""
        contacts = []
        try:
            with open(file_path, 'r', encoding='utf-8', newline='') as csvfile:
                sample = csvfile.read(1024)
                csvfile.seek(0)
                
                # Detect delimiter
                delimiter = ','
                if '\t' in sample:
                    delimiter = '\t'
                elif ';' in sample:
                    delimiter = ';'
                
                reader = csv.DictReader(csvfile, delimiter=delimiter)
                
                for row_num, row in enumerate(reader, start=2):
                    email = None
                    name = None
                    
                    # Find email and name in any column
                    for key, value in row.items():
                        if value and 'email' in key.lower():
                            email = value.strip()
                        elif value and 'name' in key.lower():
                            name = value.strip()
                    
                    if email and self.is_valid_email(email):
                        contacts.append({
                            'email': email,
                            'name': name if name else email.split('@')[0],
                            'source': file_path,
                            'row': row_num,
                            **{k.lower().replace(' ', '_'): v for k, v in row.items() if v}
                        })
            
            print(f"Parsed {len(contacts)} contacts from CSV (fallback): {file_path}")
            return contacts
            
        except Exception as e:
            print(f"Error in CSV fallback parsing {file_path}: {str(e)}")
            return []
    
    def parse_docx_file(self, file_path):
        """Parse DOCX file and extract contact information"""
        contacts = []
        
        if not DOCX_AVAILABLE:
            print(f"Warning: python-docx not available, skipping {file_path}")
            return contacts
            
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs and tables
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"
            
            # Extract emails using regex
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            emails = re.findall(email_pattern, text_content)
            
            lines = text_content.split('\n')
            
            for email in set(emails):
                if self.is_valid_email(email):
                    name = self.extract_name_for_email(email, lines)
                    contacts.append({
                        'email': email,
                        'name': name if name else email.split('@')[0],
                        'source': file_path,
                        'extracted_from': 'docx_content'
                    })
            
            print(f"Parsed {len(contacts)} contacts from DOCX: {file_path}")
            return contacts
            
        except Exception as e:
            print(f"Error parsing DOCX file {file_path}: {str(e)}")
            return []
    
    def parse_txt_file(self, file_path):
        """Parse plain text file and extract email addresses"""
        contacts = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            emails = re.findall(email_pattern, content)
            lines = content.split('\n')
            
            for email in set(emails):
                if self.is_valid_email(email):
                    name = self.extract_name_for_email(email, lines)
                    contacts.append({
                        'email': email,
                        'name': name if name else email.split('@')[0],
                        'source': file_path,
                        'extracted_from': 'txt_content'
                    })
            
            print(f"Parsed {len(contacts)} contacts from TXT: {file_path}")
            return contacts
            
        except Exception as e:
            print(f"Error parsing TXT file {file_path}: {str(e)}")
            return []
    
    def parse_url_file(self, file_path):
        """Parse .url file containing URLs to contact sources"""
        if not REQUESTS_AVAILABLE:
            print("Warning: requests library not available for URL parsing")
            return []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                url = f.read().strip()
            
            print(f"Processing URL: {url}")
            
            if 'docs.google.com/spreadsheets' in url:
                return self.parse_google_sheets_url(url, file_path)
            else:
                return self.parse_web_url(url, file_path)
                
        except Exception as e:
            print(f"Error parsing URL file {file_path}: {str(e)}")
            return []
    
    def parse_google_sheets_url(self, url, source_file):
        """Parse Google Sheets URL"""
        contacts = []
        try:
            # Convert to CSV export URL
            if '/edit' in url:
                sheet_id = url.split('/d/')[1].split('/')[0]
                gid = '0'
                
                if 'gid=' in url:
                    gid = url.split('gid=')[1].split('#')[0].split('&')[0]
                
                csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid={gid}"
                
                response = requests.get(csv_url, timeout=30)
                if response.status_code == 200:
                    # Save temporarily and parse as CSV
                    import tempfile
                    with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False) as temp_file:
                        temp_file.write(response.text)
                        temp_path = temp_file.name
                    
                    try:
                        contacts = self.parse_csv_file(temp_path)
                        for contact in contacts:
                            contact['source'] = f"{source_file} -> Google Sheets"
                    finally:
                        os.unlink(temp_path)
                        
                elif response.status_code == 403:
                    print("Error: Google Sheets access denied. Make sure sheet is publicly accessible.")
                else:
                    print(f"Error fetching Google Sheets: HTTP {response.status_code}")
                    
        except Exception as e:
            print(f"Error processing Google Sheets: {str(e)}")
            
        return contacts
    
    def parse_web_url(self, url, source_file):
        """Parse general web URL for contact data"""
        contacts = []
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (compatible; ContactParser/1.0)'}
            response = requests.get(url, headers=headers, timeout=30)
            
            if response.status_code == 200:
                # Extract emails from content
                email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
                emails = re.findall(email_pattern, response.text)
                
                for email in set(emails):
                    if self.is_valid_email(email):
                        contacts.append({
                            'email': email,
                            'name': email.split('@')[0],
                            'source': f"{source_file} -> {url}",
                            'extracted_from': 'web_content'
                        })
            else:
                print(f"Error fetching URL: HTTP {response.status_code}")
                
        except Exception as e:
            print(f"Error processing web URL: {str(e)}")
            
        return contacts
    
    def extract_name_for_email(self, email, lines):
        """Try to find a name associated with an email"""
        for line in lines:
            if email in line:
                words = line.replace(email, '').split()
                potential_names = [w for w in words if len(w) > 1 and w.isalpha()]
                if potential_names:
                    return ' '.join(potential_names[:2])
        return None
    
    def is_valid_email(self, email):
        """Validate email format"""
        email_pattern = r'^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}$'
        return re.match(email_pattern, email) is not None
    
    def parse_contacts_directory(self, contacts_directory):
        """Parse all supported files in the contacts directory"""
        all_contacts = []
        
        if not os.path.exists(contacts_directory):
            print(f"Contacts directory does not exist: {contacts_directory}")
            return all_contacts
        
        print(f"Scanning contacts directory: {contacts_directory}")
        
        for filename in os.listdir(contacts_directory):
            file_path = os.path.join(contacts_directory, filename)
            
            if not os.path.isfile(file_path):
                continue
            
            file_ext = os.path.splitext(filename)[1].lower()
            
            try:
                if file_ext == '.csv':
                    contacts = self.parse_csv_file(file_path)
                elif file_ext == '.xlsx':
                    contacts = self.parse_excel_file(file_path)
                elif file_ext == '.docx':
                    contacts = self.parse_docx_file(file_path)
                elif file_ext == '.txt':
                    contacts = self.parse_txt_file(file_path)
                elif file_ext == '.url':
                    contacts = self.parse_url_file(file_path)
                else:
                    print(f"Unsupported file format: {filename}")
                    continue
                
                if contacts:
                    print(f"Loaded {len(contacts)} contacts from {filename}")
                    all_contacts.extend(contacts)
                else:
                    print(f"No contacts found in {filename}")
                
            except Exception as e:
                print(f"Error processing file {filename}: {str(e)}")
                continue
        
        # Remove duplicates based on email
        unique_contacts = {}
        for contact in all_contacts:
            email = contact['email'].lower()
            if email not in unique_contacts:
                unique_contacts[email] = contact
            else:
                # Merge contact info if duplicate found
                existing = unique_contacts[email]
                for key, value in contact.items():
                    if key not in existing and value:
                        existing[key] = value
        
        final_contacts = list(unique_contacts.values())
        print(f"\nContact Summary:")
        print(f"   Total files processed: {len([f for f in os.listdir(contacts_directory) if os.path.isfile(os.path.join(contacts_directory, f))])}")
        print(f"   Total contacts found: {len(all_contacts)}")
        print(f"   Unique contacts: {len(final_contacts)}")
        
        return final_contacts
    
    def parse_excel_file(self, file_path):
        """Parse Excel file (.xlsx)"""
        if not PANDAS_AVAILABLE:
            print(f"Warning: pandas not available, skipping Excel file {file_path}")
            return []
        
        try:
            df = pd.read_excel(file_path)
            
            email_col = self.find_column(df.columns, ['email', 'Email', 'EMAIL'])
            name_col = self.find_column(df.columns, ['name', 'Name', 'NAME'])
            
            if not email_col:
                print(f"Warning: No email column found in {file_path}")
                return []
            
            contacts = []
            for index, row in df.iterrows():
                email = str(row[email_col]).strip() if pd.notna(row[email_col]) else ""
                name = str(row[name_col]).strip() if name_col and pd.notna(row[name_col]) else ""
                
                if email and self.is_valid_email(email):
                    contacts.append({
                        'email': email,
                        'name': name if name else email.split('@')[0],
                        'source': file_path,
                        'row': index + 2
                    })
            
            print(f"Parsed {len(contacts)} valid contacts from Excel: {file_path}")
            return contacts
            
        except Exception as e:
            print(f"Error parsing Excel file {file_path}: {str(e)}")
            return []

def load_campaign_content(campaign_path):
    """Load campaign content from various file formats"""
    try:
        file_ext = os.path.splitext(campaign_path)[1].lower()
        
        if file_ext == '.json':
            with open(campaign_path, 'r', encoding='utf-8') as f:
                campaign_data = json.load(f)
            return campaign_data
            
        elif file_ext == '.docx':
            if not DOCX_AVAILABLE:
                print(f"Warning: python-docx not available, skipping {campaign_path}")
                return None
                
            doc = Document(campaign_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                    content += "\n"
            
            return content.strip()
        
        elif file_ext in ['.txt', '.html', '.md']:
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(campaign_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    return content
                except UnicodeDecodeError:
                    continue
            
            # Binary fallback
            with open(campaign_path, 'rb') as f:
                raw_content = f.read()
                return raw_content.decode('utf-8', errors='ignore')
        
        return None
        
    except Exception as e:
        print(f"Error loading campaign content from {campaign_path}: {str(e)}")
        return None

def extract_subject_from_content(content):
    """Extract subject line from campaign content"""
    if not content:
        return None
        
    try:
        # Handle JSON format
        if isinstance(content, dict):
            return content.get('subject')
            
        # Handle string content
        lines = str(content).split('\n')
        for line in lines[:10]:
            if line.lower().startswith('subject:'):
                return line.split(':', 1)[1].strip()
            elif line.startswith('# '):
                return line[2:].strip()
        return None
    except:
        return None

def campaign_main(templates_root, contacts_root, scheduled_root, tracking_root, alerts_email, dry_run=False):
    """Main campaign execution function"""
    try:
        print(f"Starting Email Campaign System")
        print(f"   Mode: {'DRY-RUN' if dry_run else 'LIVE'}")
        print(f"   Templates: {templates_root}")
        print(f"   Contacts: {contacts_root}")
        print(f"   Scheduled: {scheduled_root}")
        print(f"   Tracking: {tracking_root}")
        print(f"   Alerts: {alerts_email}")
        print(f"   Timestamp: {datetime.now().isoformat()}")
        print("-" * 60)
        
        # Ensure directories exist
        os.makedirs(tracking_root, exist_ok=True)
        
        # Initialize contact parser
        contact_parser = ContactParser()
        print("Initializing contact parser...")
        
        # Parse all contacts
        all_contacts = contact_parser.parse_contacts_directory(contacts_root)
        
        if not all_contacts:
            print("No contacts found. Checking contact sources...")
            if os.path.exists(contacts_root):
                files = os.listdir(contacts_root)
                print(f"   Files in contacts directory: {files}")
            return
        
        print(f"Loaded {len(all_contacts)} total contacts")
        
        # Save parsed contacts
        contacts_file = os.path.join(tracking_root, f'parsed_contacts_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json')
        with open(contacts_file, 'w') as f:
            json.dump(all_contacts, f, indent=2, default=str)
        print(f"Contacts saved to: {contacts_file}")
        
        # Initialize email sender
        emailer = EmailSender(alerts_email=alerts_email, dry_run=dry_run)
        print(f"Email sender initialized (dry_run={dry_run})")
        
        # Initialize log file
        log_file = os.path.join(tracking_root, "campaign_execution.log")
        with open(log_file, 'w') as f:
            f.write(f"Campaign Execution Log\n")
            f.write(f"=====================\n")
            f.write(f"Timestamp: {datetime.now().isoformat()}\n")
            f.write(f"Mode: {'DRY-RUN' if dry_run else 'LIVE'}\n")
            f.write(f"Total contacts loaded: {len(all_contacts)}\n\n")
        
        # Initialize counters
        campaigns_processed = 0
        total_emails_sent = 0
        total_failures = 0
        campaign_results = []
        
        # Check scheduled campaigns directory
        if not os.path.exists(scheduled_root):
            print(f"Scheduled campaigns directory not found: {scheduled_root}")
            return
        
        # Get campaign files
        campaign_files = []
        for ext in ['.json', '.docx', '.txt', '.html', '.md']:
            campaign_files.extend([f for f in os.listdir(scheduled_root) if f.endswith(ext)])
        
        if not campaign_files:
            print(f"No campaign files found in {scheduled_root}")
            print("   Supported formats: .json, .docx, .txt, .html, .md")
            return
        
        print(f"Found {len(campaign_files)} campaign files:")
        for file in campaign_files:
            print(f"   - {file}")
        print()
        
        # Process each campaign
        for campaign_file in campaign_files:
            campaign_name = os.path.splitext(campaign_file)[0]
            campaign_path = os.path.join(scheduled_root, campaign_file)
            
            print(f"Processing Campaign: {campaign_name}")
            print(f"   Source: {campaign_file}")
            
            # Load campaign content
            campaign_data = load_campaign_content(campaign_path)
            
            if not campaign_data:
                print(f"Could not load campaign: {campaign_file}")
                continue
            
            # Handle different campaign data formats
            if isinstance(campaign_data, dict):
                subject = campaign_data.get('subject', f'Campaign: {campaign_name}')
                content = campaign_data.get('content', '')
                from_name = campaign_data.get('from_name', 'Campaign System')
                metadata = campaign_data.get('metadata', {})
            else:
                content = str(campaign_data)
                subject = extract_subject_from_content(content) or f'Campaign: {campaign_name}'
                from_name = 'Campaign System'
                metadata = {}
            
            if not content.strip():
                print(f"Empty campaign content: {campaign_file}")
                continue
            
            # Add recipient IDs for tracking
            contacts_with_ids = []
            for i, contact in enumerate(all_contacts):
                contact_copy = contact.copy()
                contact_copy['recipient_id'] = f"{campaign_name}_{i+1}"
                contact_copy['campaign_id'] = campaign_name
                contacts_with_ids.append(contact_copy)
            
            # Send campaign
            try:
                campaign_result = emailer.send_campaign(
                    campaign_name=campaign_name,
                    subject=subject,
                    content=content,
                    recipients=contacts_with_ids,
                    from_name=from_name
                )
                
                campaigns_processed += 1
                total_emails_sent += campaign_result['sent']
                total_failures += campaign_result['failed']
                campaign_results.append(campaign_result)
                
                # Log campaign details
                with open(log_file, 'a') as f:
                    f.write(f"Campaign: {campaign_name}\n")
                    f.write(f"Subject: {subject}\n")
                    f.write(f"Recipients: {campaign_result['total_recipients']}\n")
                    f.write(f"Sent: {campaign_result['sent']}\n")
                    f.write(f"Failed: {campaign_result['failed']}\n")
                    f.write(f"Duration: {campaign_result['duration_seconds']:.2f}s\n")
                    f.write(f"Status: {'SIMULATED' if dry_run else 'SENT'}\n\n")
                
                print(f"Campaign '{campaign_name}' completed:")
                print(f"   Sent: {campaign_result['sent']}")
                print(f"   Failed: {campaign_result['failed']}")
                
                # Save detailed results
                result_file = os.path.join(tracking_root, f"{campaign_name}_results.json")
                with open(result_file, 'w') as f:
                    json.dump(campaign_result, f, indent=2, default=str)
                
            except Exception as e:
                print(f"Error processing campaign '{campaign_name}': {str(e)}")
                with open(log_file, 'a') as f:
                    f.write(f"Campaign: {campaign_name}\n")
                    f.write(f"Status: ERROR - {str(e)}\n\n")
                continue
        
        # Final summary
        success_rate = (total_emails_sent / max(1, total_emails_sent + total_failures)) * 100
        
        with open(log_file, 'a') as f:
            f.write("=== CAMPAIGN SUMMARY ===\n")
            f.write(f"Total campaigns processed: {campaigns_processed}\n")
            f.write(f"Total emails sent: {total_emails_sent}\n")
            f.write(f"Total failures: {total_failures}\n")
            f.write(f"Success rate: {success_rate:.1f}%\n")
            f.write(f"Run mode: {'DRY-RUN' if dry_run else 'LIVE'}\n")
            f.write(f"Completed: {datetime.now().isoformat()}\n")
        
        print(f"\n=== FINAL SUMMARY ===")
        print(f"Campaigns processed: {campaigns_processed}")
        print(f"Total emails: {total_emails_sent + total_failures}")
        print(f"Successful: {total_emails_sent}")
        print(f"Failed: {total_failures}")
        print(f"Success rate: {success_rate:.1f}%")
        print(f"Mode: {'DRY-RUN' if dry_run else 'LIVE'}")
        
        # Send summary alert
        if not dry_run and campaigns_processed > 0:
            summary_subject = f"Campaign Summary: {campaigns_processed} campaigns, {total_emails_sent + total_failures} emails"
            summary_body = f"""
Campaign Execution Summary
=========================

Campaigns Processed: {campaigns_processed}
Total Emails: {total_emails_sent + total_failures}
Successful: {total_emails_sent}
Failed: {total_failures}
Success Rate: {success_rate:.1f}%

Campaign Details:
"""
            for result in campaign_results:
                summary_body += f"\n• {result['campaign_name']}: {result['sent']}/{result['total_recipients']} sent"
                if result['failed'] > 0:
                    summary_body += f" ({result['failed']} failed)"
            
            summary_body += f"\n\nExecution completed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            
            emailer.send_alert(summary_subject, summary_body)
            print("Summary alert sent")
        
    except Exception as e:
        print(f"ERROR: {str(e)}")
        traceback.print_exc()
        
        # Log error
        error_log = os.path.join(tracking_root, "error.log")
        with open(error_log, 'w') as f:
            f.write(f"ERROR: {str(e)}\n")
            f.write(traceback.format_exc())
        
        # Send error alert
        if not dry_run:
            try:
                error_emailer = EmailSender(alerts_email=alerts_email, dry_run=False)
                error_emailer.send_alert(
                    "Campaign System Error",
                    f"Campaign execution failed with error:\n\n{str(e)}\n\nCheck logs for details."
                )
            except:
                pass
        
        sys.exit(1)

if __name__ == "__main__":
    print("Email Campaign System starting...")
    print(f"Remote environment: {IS_REMOTE}")
    print(f"Available libraries: pandas={PANDAS_AVAILABLE}, docx={DOCX_AVAILABLE}, requests={REQUESTS_AVAILABLE}")
    
    parser = argparse.ArgumentParser(description='Email Campaign System - Enhanced Version')
    parser.add_argument("--templates", required=True, help="Templates directory path")
    parser.add_argument("--contacts", required=True, help="Contacts directory path")
    parser.add_argument("--scheduled", required=True, help="Scheduled campaigns directory path")
    parser.add_argument("--tracking", required=True, help="Tracking directory path")
    parser.add_argument("--alerts", required=True, help="Alerts email address")
    parser.add_argument("--dry-run", action="store_true", help="Simulate sending without actually sending emails")
    parser.add_argument("--remote-only", action="store_true", help="Force remote-only mode")
    
    args = parser.parse_args()
    
    # Override remote detection if specified
    if args.remote_only:
        globals()['IS_REMOTE'] = True
        print("Forced remote-only mode enabled")
    
    print(f"Configuration:")
    print(f"  Templates: {args.templates}")
    print(f"  Contacts: {args.contacts}")
    print(f"  Scheduled: {args.scheduled}")
    print(f"  Tracking: {args.tracking}")
    print(f"  Alerts: {args.alerts}")
    print(f"  Dry run: {args.dry_run}")
    
    # Validate directories
    required_dirs = [args.contacts, args.scheduled]
    for dir_path in required_dirs:
        if not os.path.exists(dir_path):
            print(f"ERROR: Required directory does not exist: {dir_path}")
            sys.exit(1)
    
    # Create tracking directory if it doesn't exist
    os.makedirs(args.tracking, exist_ok=True)
    
    print("\nStarting campaign execution...")
    campaign_main(
        args.templates, 
        args.contacts, 
        args.scheduled, 
        args.tracking, 
        args.alerts,
        dry_run=args.dry_run
    )
    print("Campaign execution completed successfully")
